#!/usr/bin/env python3
"""Read a .docx file and output its text content.

Usage: python scripts/read_docx.py /path/to/file.docx
"""

import sys
from pathlib import Path
from docx import Document


def read_docx(file_path: str) -> str:
    """Extract text from a .docx file."""
    doc = Document(file_path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/read_docx.py /path/to/file.docx")
        sys.exit(1)

    path = sys.argv[1]
    if not Path(path).exists():
        print(f"File not found: {path}")
        sys.exit(1)

    print(read_docx(path))
