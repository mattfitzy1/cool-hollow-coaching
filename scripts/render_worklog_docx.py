#!/usr/bin/env python3
"""
Render a work log markdown file to a styled .docx.

Format matches the previous hand-built work logs: dark green Arial headings,
bulleted lists, tables for meetings, light-green callouts for observations.

Usage (CLI, for backfill):
    python scripts/render_worklog_docx.py outputs/Work_Log_16_April_2026.md
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Brand palette (Cool Hollow Coaching house style — updated 1 June 2026, was Material green)
DARK_GREEN = RGBColor(0x1F, 0x4E, 0x4A)   # brand teal (kept name to avoid touching call sites)
LIGHT_GREEN_HEX = "ECEFE7"                # sage-faint callout
GREY = RGBColor(0x66, 0x66, 0x66)
LOGO_PATH = Path(__file__).resolve().parent.parent / "outputs" / "brand" / "cool-hollow-logo" / "doc-header.png"


def _add_logo_header(doc):
    """Insert the Cool Hollow Coaching logo + bronze rule at the top of the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    if LOGO_PATH.exists():
        run.add_picture(str(LOGO_PATH), width=Inches(2.0))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "C89968")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DARK_GREEN
    r.font.name = "Arial"


def _add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = DARK_GREEN
    r.font.name = "Arial"


def _add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK_GREEN
    r.font.name = "Arial"


def _apply_inline_runs(paragraph, text):
    """Turn **bold**, *italic*, and plain text into runs on a paragraph."""
    # Split on bold first, then italic within each piece.
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
    pieces = pattern.split(text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) >= 4:
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) >= 2:
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(piece)
        run.font.size = Pt(11)
        run.font.name = "Arial"


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    _apply_inline_runs(p, text)


def _add_sub_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet 2")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    _apply_inline_runs(p, text)


def _add_paragraph(doc, text, italic=False, grey=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    _apply_inline_runs(p, text)
    if italic or grey:
        for run in p.runs:
            if italic:
                run.italic = True
            if grey:
                run.font.color.rgb = GREY
    return p


def _add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _apply_inline_runs(p, text)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), LIGHT_GREEN_HEX)
    pPr.append(shd)


def _add_footer(doc, text):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc, header_cells, rows):
    table = doc.add_table(rows=1, cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    header_row = table.rows[0].cells
    for i, h in enumerate(header_cells):
        cell = header_row[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "1B5E20")
    for row_cells in rows:
        row = table.add_row().cells
        for i, content in enumerate(row_cells):
            if i >= len(row):
                break
            cell = row[i]
            cell.text = ""
            p = cell.paragraphs[0]
            _apply_inline_runs(p, content)


def _parse_table(lines, start_idx):
    """Starting at a pipe-line, consume the table and return (rows, new_idx).

    First row is header, second is the separator (|---|---|), rest are data.
    """
    header_line = lines[start_idx]
    header = [c.strip() for c in header_line.strip().strip("|").split("|")]
    idx = start_idx + 1
    # Skip separator
    if idx < len(lines) and re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[idx]):
        idx += 1
    data = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        data.append(row)
        idx += 1
    return header, data, idx


def render_worklog_to_docx(md_text, output_path, footer_text=None):
    """Render a work log markdown string to a styled .docx file.

    footer_text: optional footer string. Only set this for internal work logs.
    Leave as None for client-facing or general documents.
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    _add_logo_header(doc)

    lines = md_text.splitlines()
    i = 0
    current_h2 = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            _add_h1(doc, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("## "):
            current_h2 = stripped[3:].strip()
            _add_h2(doc, current_h2)
            i += 1
            continue

        if stripped.startswith("### "):
            _add_h3(doc, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("|"):
            header, rows, i = _parse_table(lines, i)
            _add_table(doc, header, rows)
            continue

        # Sub-bullets (4+ spaces indent) before top-level bullet check
        if re.match(r"^\s{2,}[-*] ", line):
            text = re.sub(r"^\s+[-*]\s+", "", line)
            _add_sub_bullet(doc, text)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            _add_bullet(doc, stripped[2:])
            i += 1
            continue

        # Treat Observations / Next Steps paragraphs as callouts for emphasis
        if current_h2 and current_h2.lower().startswith("observations"):
            _add_callout(doc, stripped)
            i += 1
            continue

        _add_paragraph(doc, stripped)
        i += 1

    if footer_text:
        _add_footer(doc, footer_text)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python scripts/render_worklog_docx.py <path-to-md>")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"File not found: {md_path}")
        sys.exit(1)
    md_text = md_path.read_text()
    out_path = md_path.with_suffix(".docx")
    render_worklog_to_docx(md_text, out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
