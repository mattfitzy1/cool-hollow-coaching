"""Build the AIOS Setup Guide (.docx) using the brand house style.

A name-agnostic guide for a DIY install. Writes the branded .docx to outputs/
(and Desktop via brand_docx.save).
"""
from __future__ import annotations

from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from brand_docx import (
    new_brand_doc, add_logo_header, heading, body, bullet, save,
    TEAL, DEEP_TEAL, BRONZE, BODY as BODY_COLOR,
)

MONO = "Courier New"


def command(doc, text: str):
    """A monospaced command line on a light shaded background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(10)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F1F4F2")  # very light sage-grey
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(10.5)
    run.font.color.rgb = DEEP_TEAL
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), MONO)
    return p


def step(doc, text: str):
    """A bold bronze step label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = TEAL
    return p


def numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_COLOR
    return p


doc = new_brand_doc()
add_logo_header(
    doc,
    title="Your AIOS - DIY Setup Guide",
    subtitle="Everything you need to set up and run your AI Operating System yourself",
)

# --- Intro ---
body(doc, "Welcome. You are about to set up your own AI Operating System, an AI layer wrapped around your business that learns how you work and takes the manual, repetitive parts off your plate, one piece at a time.")
body(doc, "You do not need to be technical. If you can describe something in plain English, your AIOS can do it. The whole thing is built to be talked to, not typed at. The first setup takes about 30 to 45 minutes, most of which is your AIOS interviewing you about your business so it understands your world.")
body(doc, "Follow the steps in order. Nothing here can break anything, so feel free to explore and ask your AIOS questions about itself as you go.")

# --- The five principles ---
heading(doc, "Five principles to keep in mind")
bullet(doc, "Just ask. If you can describe it in plain English, it can probably be built. Do not self-censor. Ask for the impossible.")
bullet(doc, "Talk, do not type. Speak your answers and let the AIOS format them. It is faster and it sounds more like you.")
bullet(doc, "Layers, not leaps. Add one capability at a time. Each one is useful on its own. You get more confident without trying.")
bullet(doc, "Build for scale and security. Your data stays on your machine. A human stays in the loop by default. Plan before you build.")
bullet(doc, "Borrow before you build. There is a library of ready-made modules. Check there before building anything from scratch.")

# --- Before you start ---
heading(doc, "Before you start")
body(doc, "Three things to have ready:")
bullet(doc, "A Claude subscription (Claude Pro or Max). This is what powers everything. If you do not have one, get it at claude.ai. Max is recommended if you plan to use it heavily.")
bullet(doc, "A free GitHub account at github.com. Send us your GitHub username and we will give you access to the private workspace we have set up for you.")
bullet(doc, "A computer (these steps are written for a Mac; Windows works too with minor differences). About 30 to 45 minutes for the first sitting.")

# --- Step 1 ---
heading(doc, "Step 1 - Install Claude Code")
body(doc, "Claude Code is the app that runs your AIOS. The desktop app is the simplest way to start.")
numbered(doc, "Go to claude.ai/code and download the desktop app for your computer.")
numbered(doc, "Install it and sign in with the same account as your Claude subscription.")
body(doc, "If you prefer working in a terminal you can install the command-line version instead, but the desktop app needs no technical setup.")

# --- Step 2 ---
heading(doc, "Step 2 - Get your workspace onto your computer")
body(doc, "We have created a private workspace for you on GitHub (named after you, for example yourname-aios). You will get an email invitation to it. The simplest way to get it onto your machine, with no terminal needed:")
numbered(doc, "Accept the GitHub invitation from the email we send you.")
numbered(doc, "Open your workspace page on GitHub.")
numbered(doc, "Click the green Code button, then Download ZIP.")
numbered(doc, "Unzip it somewhere sensible, like a folder called AIOS inside your Documents.")
body(doc, "If you are comfortable with git you can clone it instead, which makes future updates cleaner. Either way works to start.")

# --- Step 3 ---
heading(doc, "Step 3 - Open it and turn off the prompts")
numbered(doc, "Open the Claude Code desktop app.")
numbered(doc, "Open the workspace folder you just unzipped.")
numbered(doc, "One-time setting so your AIOS does not ask permission for every small action: go to Settings, then Claude Code, and turn on Allow bypass permissions mode. This makes it feel like a conversation rather than a form. You only do this once and it sticks.")

# --- Step 4 ---
heading(doc, "Step 4 - Build your context layer")
body(doc, "This is the important one. Your AIOS starts as a blank template. The first thing you do is teach it about you and your business. In Claude Code, type this and send it:")
command(doc, "/install module-installs/context-os-v1")
body(doc, "Your AIOS will interview you. It asks about your business, your role, what you are trying to achieve right now, and the numbers that matter. Answer in plain English. When it is done, it has written your context files and it understands your world.")
body(doc, "Tip: when it asks a question, hold your dictation key and just talk for a minute. It is faster than typing and the AIOS keeps your natural voice.")

# --- Step 5 ---
heading(doc, "Step 5 - Start every session with a quick catch-up")
body(doc, "Once your context is built, begin each working session by typing:")
command(doc, "/prime")
body(doc, "This loads everything your AIOS knows about you and tells you where things stand. Think of it as your AIOS reading itself in before you start.")

# --- Step 6 ---
heading(doc, "Step 6 - Add capabilities, one layer at a time")
body(doc, "Your AIOS now knows who you are. From here you add capabilities one at a time. Do not install everything at once. Pick the single thing that would save you the most time this week and start there.")
body(doc, "There are ready-made modules in the module-installs folder of your workspace. To install one, type:")
command(doc, "/install module-installs/the-module-name")
body(doc, "A few worth looking at early:")
bullet(doc, "A daily brief that pulls your day together and sends it to you each morning.")
bullet(doc, "A task and project tracker so nothing slips.")
bullet(doc, "A research assistant that reads the web, papers, and video for you.")
bullet(doc, "A content helper for drafting and scheduling posts.")
body(doc, "Not sure which to pick? Just ask your AIOS: tell it what eats most of your time each week and ask which module would help most.")

# --- The five layers ---
heading(doc, "The bigger picture: five layers")
body(doc, "Your AIOS is built up in five layers. You do not need to think about this to use it, but it helps to know the road ahead.")
bullet(doc, "Context. Your AI understands your business: strategy, team, processes, history.")
bullet(doc, "Data. Your AI sees your numbers in real time, pulled from your actual sources.")
bullet(doc, "Intelligence. Your AI watches what matters and gives you a daily brief.")
bullet(doc, "Automate. You score each recurring task and automate them away one by one.")
bullet(doc, "Build. The time you get back goes into growth, new ideas, or simply your life.")

# --- Getting the most out of it ---
heading(doc, "Getting the most out of it")
bullet(doc, "Talk to it like a sharp colleague. Explain what you want in plain words. Ask follow-up questions.")
bullet(doc, "Ask for more than you think is possible. The answer is often yes.")
bullet(doc, "For anything bigger, ask it to plan first, then build. You stay in control of what happens.")
bullet(doc, "Your data stays on your machine. Nothing is shared anywhere unless you ask it to.")
bullet(doc, "Keep your context current. When something changes in your business, tell your AIOS so it stays accurate.")

# --- Support ---
heading(doc, "If you get stuck")
body(doc, "Send us a screenshot of what you are seeing and we will walk you through it. Your DIY setup comes with light mentoring, so you are not on your own. The fastest way to learn is to use it for real work from day one.")
body(doc, "Welcome to your AIOS.")

out = save(doc, "outputs/AIOS_Setup_Guide.docx")
print(f"Saved: {out}")
