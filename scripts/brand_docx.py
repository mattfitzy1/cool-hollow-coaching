"""Cool Hollow Coaching document house style for python-docx deliverables.

THE canonical brand helper. Every .docx deliverable we produce should:
  1. call add_logo_header(doc, title=...) first — puts the logo + bronze rule at the top, and
  2. use heading()/body()/bullet() so brand colours stay consistent.

This replaces the old ad-hoc `#1B5E20` Material green with the real brand teal.
Fonts: headings in Georgia (serif, evokes the Fraunces wordmark — universally
installed so nothing substitutes), body in Arial. The brand identity is carried
by the logo image + the teal/bronze palette, which render identically everywhere.

Quick start:
    from brand_docx import new_brand_doc, add_logo_header, heading, body, bullet, save
    doc = new_brand_doc()
    add_logo_header(doc, title="Meeting Prep — Acme Co", subtitle="1 June 2026")
    heading(doc, "Context")
    body(doc, "One or two lines of setup.")
    bullet(doc, "First point")
    save(doc, "outputs/example.docx")   # also copies to Desktop
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Brand palette (matches the live site + slide decks) ---
TEAL      = RGBColor(0x1F, 0x4E, 0x4A)   # primary heading / brand signal
DEEP_TEAL = RGBColor(0x0E, 0x2B, 0x2A)   # darkest, near-black ink
SAGE      = RGBColor(0xA8, 0xB5, 0xA0)
BRONZE    = RGBColor(0xC8, 0x99, 0x68)   # accent rule, emphasis
INK       = RGBColor(0x0E, 0x2B, 0x2A)
BODY      = RGBColor(0x4A, 0x59, 0x57)
MUTED     = RGBColor(0x7A, 0x85, 0x83)

HEADING_FONT = "Georgia"
BODY_FONT    = "Arial"

REPO = Path(__file__).resolve().parent.parent
LOGO_PATH = REPO / "outputs" / "brand" / "cool-hollow-logo" / "doc-header.png"
DESKTOP = Path.home() / "Desktop"


def _set_font(run, name=BODY_FONT, size=11, color=BODY, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    # ensure East-Asian/complex fallbacks also use the font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)


def _add_bottom_rule(paragraph, hex_color="C89968", size=14):
    """Draw a coloured rule along the bottom border of a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))      # eighths of a point
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def new_brand_doc(margin_in: float = 1.0) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(margin_in)
        section.bottom_margin = Inches(margin_in)
        section.left_margin = Inches(margin_in)
        section.right_margin = Inches(margin_in)
    # default Normal style → brand body
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    return doc


def add_logo_header(doc: Document, title: str | None = None,
                    subtitle: str | None = None, logo_width_in: float = 2.0) -> Document:
    """Insert the Cool Hollow Coaching logo at the top, an optional title + subtitle, and a bronze rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    if LOGO_PATH.exists():
        run.add_picture(str(LOGO_PATH), width=Inches(logo_width_in))
    _add_bottom_rule(p, "C89968", size=14)

    if title:
        t = doc.add_paragraph()
        t.paragraph_format.space_before = Pt(12)
        t.paragraph_format.space_after = Pt(2)
        _set_font(t.add_run(title), name=HEADING_FONT, size=22, color=TEAL, bold=True)
    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(10)
        _set_font(s.add_run(subtitle), name=BODY_FONT, size=11.5, color=MUTED)
    return doc


def heading(doc: Document, text: str, level: int = 2) -> None:
    sizes = {1: 18, 2: 15, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(3)
    _set_font(p.add_run(text), name=HEADING_FONT, size=sizes.get(level, 13),
              color=TEAL, bold=True)


def body(doc: Document, text: str, size: float = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    _set_font(p.add_run(text), name=BODY_FONT, size=size, color=BODY)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(text), name=BODY_FONT, size=11, color=BODY)


def save(doc: Document, path: str | Path, to_desktop: bool = True) -> Path:
    path = Path(path).resolve()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    if to_desktop:
        try:
            shutil.copy2(path, DESKTOP / path.name)
        except Exception:
            pass
    return path


if __name__ == "__main__":
    # Demo / smoke test
    doc = new_brand_doc()
    add_logo_header(doc, title="Sample Deliverable", subtitle="Brand house-style demo · 1 June 2026")
    heading(doc, "What this is", 2)
    body(doc, "Every deliverable we produce now carries the Cool Hollow Coaching logo at the top, "
              "a bronze rule, and the brand palette. This is the canonical house style.")
    heading(doc, "Key points", 2)
    bullet(doc, "Logo + bronze rule header via add_logo_header().")
    bullet(doc, "Headings in brand teal; body in a clean, readable grey.")
    bullet(doc, "One import, used by every docx generator going forward.")
    out = save(doc, REPO / "outputs" / "brand" / "brand-docx-demo.docx", to_desktop=True)
    print(f"Wrote {out}")
